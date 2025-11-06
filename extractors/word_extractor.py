from docx import Document
from typing import List

class WordExtractor:
    def extract_text(self, file) -> str:
        """Extract text from Word document"""
        doc = Document(file)
        text = ""
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n\n"
        
        # Extract tables
        if doc.tables:
            text += "\n--- Tables Found ---\n"
            for i, table in enumerate(doc.tables):
                text += f"\nTable {i+1}:\n"
                text += self.extract_table(table)
                text += "\n"
                
        return text
    
    def extract_table(self, table) -> str:
        """Extract data from a Word table"""
        table_text = ""
        
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_text += " | ".join(row_data) + "\n"
            
        return table_text